import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        # print(para.text)
        fullText.append(para.text)
    return '\n'.join(fullText), fullText

def get_dict(filename):
    doc, dlist = getText(filename)
    # print(doc)
    # print(dlist[0])
    # print(dlist)
    dict = {}
    # print(type(dict))
    # print(len(dlist[5]))
    flag = 0   # for checking whether the previous was the question or not
    ques = ''
    for eachl in dlist:
        if len(eachl) == 0:
            continue
        # print(eachl[0])
        if eachl[0] == 'Q':
            ques = eachl[0:2]
            dict[ques] = eachl.split(']')[2]
            flag = 1
        elif flag == 1:
            prevstr = dict[ques]
            curstr = eachl
            joinstr = prevstr + ' ' + curstr
            dict[ques] = joinstr
    return dict

if __name__ == "__main__":
    dict = get_dict("T1.docx")
    for key,ele in dict.items():
        print(key, ' : ', ele)



# document is returning object of type Document
# paragraph is returning a list of objects of type paragraph
# text is converting that paragraphs into text

# doc = docx.Document("T1.docx")
# print(doc)
# print(doc.paragraphs)
# print(doc.paragraphs[0].text)


# f = open("T1.docx", "r")
# print(f.readline())
# f.close()